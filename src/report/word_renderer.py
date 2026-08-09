from __future__ import annotations

import os
from dataclasses import dataclass, field
from datetime import date, datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Sequence

from docx import Document as DocxDocumentFactory
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor

from .models import ReportCitationDraft, ReportSectionDraft, ReportSectionStatus
from .pdf_renderer import normalize_pdf_text


DEFAULT_WORD_FONT_NAME = "Malgun Gothic"
DEFAULT_WORD_TITLE = "일일 산업 동향 보고서"
PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_WORD_LOGO_PATH = PROJECT_ROOT / "assets" / "brand" / "mySUNI.png"
WORD_LOGO_ENV = "MYWIKI_WORD_LOGO_PATH"

K_EXECUTIVE_SUMMARY = "오늘의 핵심 요약"
K_MAJOR_ISSUES = "주요 Issue"
K_CATEGORY_SUMMARY = "카테고리별 정리"
K_OVERALL_IMPLICATIONS = "종합 시사점"
K_SOURCES = "전체 출처"
K_CURRENT_STATUS = "현재 상황"
K_KEY_FACTS = "핵심 사실"
K_CONTEXT = "과거 맥락"
K_SK_IMPACT = "SK하이닉스 영향"
K_WATCH_POINTS = "다음 확인 사항"
K_NO_INFORMATION = "정보 없음"
K_NO_MAJOR_TREND = "주요 동향 없음"
K_NEWS = "NEWS"
K_MYWIKI = "MYWIKI"
K_INTERNAL_WIKI = "내부 Wiki"
K_LINK = "링크"
CATEGORY_ORDER = ("제품·기술", "경쟁사", "고객·수요산업", "공급망·생산", "정책·규제", "시장·경영")

NAVY = "0B2545"
ACCENT = "EA580C"
DARK_GRAY = "374151"
MEDIUM_GRAY = "6B7280"
LIGHT_GRAY = "F3F4F6"
BORDER_GRAY = "D1D5DB"
LINK_BLUE = "2563EB"
WHITE = "FFFFFF"


@dataclass(frozen=True)
class WordLayout:
    page_width_mm: float = 210
    page_height_mm: float = 297
    top_margin_mm: float = 19
    bottom_margin_mm: float = 17
    left_margin_mm: float = 19
    right_margin_mm: float = 19
    header_distance_mm: float = 8
    footer_distance_mm: float = 8
    title_font_size: float = 22
    subtitle_font_size: float = 10.5
    section_title_font_size: float = 15.5
    issue_label_font_size: float = 9
    issue_title_font_size: float = 14
    subheading_font_size: float = 11
    body_font_size: float = 10
    bullet_font_size: float = 10
    metadata_font_size: float = 8.8
    source_font_size: float = 8.8
    footer_font_size: float = 7.8
    body_line_spacing: float = 1.35
    bullet_line_spacing: float = 1.32
    source_line_spacing: float = 1.25
    max_evidences_per_section: int = 6
    max_executive_items: int = 5
    max_category_items: int = 3
    header_logo_height_mm: float = 7
    header_brand_font_size: float = 10
    header_date_font_size: float = 9

    @property
    def usable_width_mm(self) -> float:
        return self.page_width_mm - self.left_margin_mm - self.right_margin_mm


@dataclass(frozen=True)
class WordEvidenceLine:
    document_version_id: str
    citation_order: int
    evidence_text: str
    relevance_score: Optional[float] = None
    document_title: str | None = None
    source_name: str | None = None
    published_at: str | None = None
    source_url: str | None = None


@dataclass(frozen=True)
class WordExecutiveSummaryLine:
    title: str
    summary: str
    category: str = ""
    importance_score: Optional[int] = None
    impact_direction: Optional[str] = None
    time_horizon: Optional[str] = None


@dataclass(frozen=True)
class WordCategorySummary:
    category: str
    items: tuple[str, ...] = ()


@dataclass(frozen=True)
class WordOverallImplications:
    opportunities: tuple[str, ...] = ()
    risks: tuple[str, ...] = ()
    monitoring_points: tuple[str, ...] = ()


@dataclass(frozen=True)
class WordSourceLine:
    source_type: str = ""
    source_name: str = ""
    title: str = ""
    published_at: str = ""
    url: str | None = None


@dataclass(frozen=True)
class WordSection:
    category: str
    title: str
    status: str
    importance_score: Optional[int]
    impact_direction: Optional[str]
    time_horizon: Optional[str]
    current_summary: Optional[str]
    key_facts: tuple[str, ...] = ()
    historical_context: tuple[str, ...] = ()
    implications: tuple[str, ...] = ()
    watch_points: tuple[str, ...] = ()
    evidences: tuple[WordEvidenceLine, ...] = ()


@dataclass(frozen=True)
class WordReportDocument:
    title: str
    subtitle: str
    generated_at: str
    report_date: str
    version: int
    layout: WordLayout = field(default_factory=WordLayout)
    sections: tuple[WordSection, ...] = ()
    executive_items: tuple[WordExecutiveSummaryLine, ...] = ()
    category_summaries: tuple[WordCategorySummary, ...] = ()
    overall_implications: WordOverallImplications | None = None
    source_rows: tuple[WordSourceLine, ...] = ()


DEFAULT_DAILY_REPORT_WORD_LAYOUT = WordLayout()


def build_daily_report_word_document(
    *,
    report_key: str,
    version: int,
    sections: list[ReportSectionDraft],
    generated_at: Optional[str] = None,
    report_date: date | str | None = None,
    title: Optional[str] = None,
    layout: WordLayout = DEFAULT_DAILY_REPORT_WORD_LAYOUT,
    executive_summaries: Sequence[Any] | None = None,
    issue_summary_rows: Sequence[Any] | None = None,
    category_groups: Sequence[Any] | None = None,
    overall_implications: Any | None = None,
    news_sources: Sequence[Any] | None = None,
    wiki_sources: Sequence[Any] | None = None,
) -> WordReportDocument:
    completed_sections = [
        _to_word_section(section, max_evidences=layout.max_evidences_per_section)
        for section in sections
        if _status_value(section.status) == ReportSectionStatus.COMPLETED.value
    ]
    normalized_generated_at = normalize_pdf_text(generated_at or _utc_now_iso())
    normalized_report_date = normalize_pdf_text(_resolve_report_date_text(report_date=report_date, generated_at=normalized_generated_at))
    return WordReportDocument(
        title=normalize_pdf_text(title or DEFAULT_WORD_TITLE),
        subtitle=normalize_pdf_text(report_key),
        generated_at=normalized_generated_at,
        report_date=normalized_report_date,
        version=version,
        layout=layout,
        sections=tuple(completed_sections),
        executive_items=_build_executive_items(completed_sections, executive_summaries, issue_summary_rows, layout.max_executive_items),
        category_summaries=_build_category_summaries(completed_sections, category_groups, layout.max_category_items),
        overall_implications=_to_word_overall_implications(overall_implications),
        source_rows=_build_source_rows(completed_sections, news_sources, wiki_sources),
    )


def build_daily_report_word_filename(*, report_key: str, version: int) -> str:
    normalized_key = normalize_pdf_text(report_key).strip().replace(" ", "-")
    return f"{normalized_key}-v{version}.docx"


def render_daily_report_word(document: WordReportDocument) -> bytes:
    normalized_document = _normalize_document(document)
    doc = DocxDocumentFactory()
    _configure_page_layout(doc, normalized_document.layout)
    _configure_styles(doc, normalized_document.layout)
    _configure_header(doc, normalized_document)
    _configure_footer(doc, normalized_document.layout)
    doc.core_properties.title = normalized_document.title
    doc.core_properties.author = "myWiki"

    _render_cover_summary(doc, normalized_document)
    if normalized_document.sections:
        doc.add_page_break()
        _add_section_title(doc, K_MAJOR_ISSUES)
        for index, section in enumerate(normalized_document.sections, start=1):
            if index > 1:
                _add_section_gap(doc)
            _render_issue_section(doc, index, section)
    else:
        _add_body_paragraph(doc, "완료된 섹션이 없습니다.")

    if normalized_document.category_summaries:
        doc.add_page_break()
        _render_category_section(doc, normalized_document)
    if normalized_document.overall_implications is not None:
        doc.add_page_break()
        _render_overall_implications(doc, normalized_document.overall_implications, normalized_document.layout)
    if normalized_document.source_rows:
        doc.add_page_break()
        _render_sources_section(doc, normalized_document.source_rows, normalized_document.layout)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _resolve_report_date_text(*, report_date: date | str | None, generated_at: str) -> str:
    if isinstance(report_date, date):
        return report_date.isoformat()
    if isinstance(report_date, str) and report_date.strip():
        raw = report_date.strip()
        if "T" in raw:
            try:
                return datetime.fromisoformat(raw.replace("Z", "+00:00")).date().isoformat()
            except ValueError:
                return raw.split("T", 1)[0]
        return raw
    try:
        return datetime.fromisoformat(generated_at.replace("Z", "+00:00")).date().isoformat()
    except ValueError:
        return generated_at.split("T", 1)[0]


def _normalize_document(document: WordReportDocument) -> WordReportDocument:
    return WordReportDocument(
        title=normalize_pdf_text(document.title),
        subtitle=normalize_pdf_text(document.subtitle),
        generated_at=normalize_pdf_text(document.generated_at),
        report_date=normalize_pdf_text(document.report_date),
        version=document.version,
        layout=document.layout,
        sections=tuple(_normalize_section(section) for section in document.sections),
        executive_items=tuple(WordExecutiveSummaryLine(normalize_pdf_text(i.title), normalize_pdf_text(i.summary), normalize_pdf_text(i.category), i.importance_score, normalize_pdf_text(i.impact_direction) if i.impact_direction else None, normalize_pdf_text(i.time_horizon) if i.time_horizon else None) for i in document.executive_items),
        category_summaries=tuple(WordCategorySummary(normalize_pdf_text(s.category), tuple(normalize_pdf_text(i) for i in s.items if _has_value(i))) for s in document.category_summaries),
        overall_implications=_normalize_overall(document.overall_implications),
        source_rows=tuple(_normalize_source(row) for row in document.source_rows),
    )


def _normalize_section(section: WordSection) -> WordSection:
    return WordSection(
        category=normalize_pdf_text(section.category),
        title=normalize_pdf_text(section.title),
        status=normalize_pdf_text(section.status),
        importance_score=section.importance_score,
        impact_direction=normalize_pdf_text(section.impact_direction) if section.impact_direction else None,
        time_horizon=normalize_pdf_text(section.time_horizon) if section.time_horizon else None,
        current_summary=normalize_pdf_text(section.current_summary) if section.current_summary else None,
        key_facts=tuple(normalize_pdf_text(i) for i in section.key_facts if _has_value(i)),
        historical_context=tuple(normalize_pdf_text(i) for i in section.historical_context if _has_value(i)),
        implications=tuple(normalize_pdf_text(i) for i in section.implications if _has_value(i)),
        watch_points=tuple(normalize_pdf_text(i) for i in section.watch_points if _has_value(i)),
        evidences=tuple(WordEvidenceLine(normalize_pdf_text(e.document_version_id), e.citation_order, normalize_pdf_text(e.evidence_text or ""), e.relevance_score, normalize_pdf_text(e.document_title) if e.document_title else None, normalize_pdf_text(e.source_name) if e.source_name else None, normalize_pdf_text(e.published_at) if e.published_at else None, normalize_pdf_text(e.source_url) if e.source_url else None) for e in section.evidences),
    )


def _normalize_overall(overall: WordOverallImplications | None) -> WordOverallImplications | None:
    if overall is None:
        return None
    return WordOverallImplications(
        opportunities=tuple(normalize_pdf_text(i) for i in overall.opportunities if _has_value(i)),
        risks=tuple(normalize_pdf_text(i) for i in overall.risks if _has_value(i)),
        monitoring_points=tuple(normalize_pdf_text(i) for i in overall.monitoring_points if _has_value(i)),
    )


def _normalize_source(source: WordSourceLine) -> WordSourceLine:
    return WordSourceLine(normalize_pdf_text(source.source_type), normalize_pdf_text(source.source_name), normalize_pdf_text(source.title), _format_source_date(source.published_at), normalize_pdf_text(source.url) if source.url else None)


def _to_word_section(section: ReportSectionDraft, *, max_evidences: int) -> WordSection:
    return WordSection(
        category=section.category.value,
        title=section.title,
        status=_status_value(section.status),
        importance_score=section.importance_score,
        impact_direction=section.impact_direction.value if section.impact_direction else None,
        time_horizon=section.time_horizon.value if section.time_horizon else None,
        current_summary=section.current_summary,
        key_facts=tuple(section.key_facts),
        historical_context=tuple(section.historical_context),
        implications=tuple(section.implications),
        watch_points=tuple(section.watch_points),
        evidences=tuple(_to_word_evidence(citation) for citation in section.news_citations[:max_evidences]),
    )


def _to_word_evidence(citation: ReportCitationDraft) -> WordEvidenceLine:
    return WordEvidenceLine(
        document_version_id=citation.document_version_id,
        citation_order=citation.citation_order,
        evidence_text=(citation.evidence_text or "").strip() or "Citation reference",
        relevance_score=citation.relevance_score,
        document_title=citation.document_title,
        source_name=citation.source_name,
        published_at=citation.published_at,
        source_url=citation.source_url,
    )


def _build_executive_items(completed_sections: Sequence[WordSection], executive_summaries: Sequence[Any] | None, issue_summary_rows: Sequence[Any] | None, max_items: int) -> tuple[WordExecutiveSummaryLine, ...]:
    category_by_issue = {str(getattr(row, "issue_key", "")): _enum_text(getattr(row, "category", "")) for row in (issue_summary_rows or ())}
    items: list[WordExecutiveSummaryLine] = []
    for item in list(executive_summaries or ())[:max_items]:
        issue_key = str(getattr(item, "issue_key", ""))
        items.append(WordExecutiveSummaryLine(
            title=_text_or_dash(getattr(item, "title", "")),
            summary=_text_or_dash(getattr(item, "summary", "")),
            category=category_by_issue.get(issue_key, ""),
            importance_score=getattr(item, "importance_score", None),
            impact_direction=_enum_text(getattr(item, "impact_direction", None)),
            time_horizon=_enum_text(getattr(item, "time_horizon", None)),
        ))
    if items:
        return tuple(items)
    return tuple(WordExecutiveSummaryLine(s.title, s.current_summary or "", s.category, s.importance_score, s.impact_direction, s.time_horizon) for s in completed_sections[:max_items] if s.title or s.current_summary)


def _build_category_summaries(completed_sections: Sequence[WordSection], category_groups: Sequence[Any] | None, max_items: int) -> tuple[WordCategorySummary, ...]:
    groups: dict[str, list[str]] = {category: [] for category in CATEGORY_ORDER}
    if category_groups:
        for group in category_groups:
            category = _enum_text(getattr(group, "category", ""))
            if not category:
                continue
            groups.setdefault(category, [])
            for section in list(getattr(group, "sections", []) or [])[:max_items]:
                groups[category].append(_join_title_summary(_text_or_dash(getattr(section, "title", "")), _clean_optional(getattr(section, "current_summary", None))))
    else:
        for section in completed_sections:
            groups.setdefault(section.category, [])
            if len(groups[section.category]) < max_items:
                groups[section.category].append(_join_title_summary(section.title, section.current_summary))
    ordered = list(CATEGORY_ORDER) + [category for category in groups if category not in CATEGORY_ORDER]
    return tuple(WordCategorySummary(category, tuple(groups.get(category, []))) for category in ordered)


def _to_word_overall_implications(overall: Any | None) -> WordOverallImplications | None:
    if overall is None:
        return None
    return WordOverallImplications(tuple(getattr(overall, "opportunities", []) or ()), tuple(getattr(overall, "risks", []) or ()), tuple(getattr(overall, "monitoring_points", []) or ()))


def _build_source_rows(completed_sections: Sequence[WordSection], news_sources: Sequence[Any] | None, wiki_sources: Sequence[Any] | None) -> tuple[WordSourceLine, ...]:
    rows: list[WordSourceLine] = []
    for source in news_sources or ():
        rows.append(WordSourceLine(K_NEWS, _clean_optional(getattr(source, "source_name", None)) or _clean_optional(getattr(source, "document_version_id", None)), _clean_optional(getattr(source, "document_title", None)), _format_source_date(getattr(source, "published_at", None)), _clean_optional(getattr(source, "source_url", None)) or None))
    for source in wiki_sources or ():
        rows.append(WordSourceLine(K_MYWIKI, K_INTERNAL_WIKI, _clean_optional(getattr(source, "wiki_title", None)) or _clean_optional(getattr(source, "wiki_page_id", None)), "", None))
    if rows:
        return tuple(_dedupe_sources(rows))
    fallback: list[WordSourceLine] = []
    for section in completed_sections:
        for evidence in section.evidences:
            fallback.append(WordSourceLine("", evidence.source_name or evidence.document_version_id, evidence.document_title or evidence.evidence_text, _format_source_date(evidence.published_at), evidence.source_url))
    return tuple(_dedupe_sources(fallback))


def _dedupe_sources(rows: Sequence[WordSourceLine]) -> list[WordSourceLine]:
    deduped: list[WordSourceLine] = []
    seen: set[tuple[str, str, str, str, str]] = set()
    for row in rows:
        key = (row.source_type, row.source_name, row.title, row.published_at, row.url or "")
        if key not in seen:
            seen.add(key)
            deduped.append(row)
    return deduped


def _configure_page_layout(doc, layout: WordLayout) -> None:
    for section in doc.sections:
        section.page_width = Mm(layout.page_width_mm)
        section.page_height = Mm(layout.page_height_mm)
        section.top_margin = Mm(layout.top_margin_mm)
        section.bottom_margin = Mm(layout.bottom_margin_mm)
        section.left_margin = Mm(layout.left_margin_mm)
        section.right_margin = Mm(layout.right_margin_mm)
        section.header_distance = Mm(layout.header_distance_mm)
        section.footer_distance = Mm(layout.footer_distance_mm)
        section.start_type = WD_SECTION_START.NEW_PAGE


def _configure_styles(doc, layout: WordLayout) -> None:
    normal = doc.styles["Normal"]
    _apply_style_font(normal, size=layout.body_font_size, color=DARK_GRAY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = layout.body_line_spacing
    normal.paragraph_format.widow_control = True
    _ensure_paragraph_style(doc, "MyWikiReportTitle", size=layout.title_font_size, color=NAVY, bold=True, before=0, after=3, line_spacing=1.18)
    _ensure_paragraph_style(doc, "MyWikiTitle", alias_of="MyWikiReportTitle")
    _ensure_paragraph_style(doc, "MyWikiSubtitle", size=layout.subtitle_font_size, color=MEDIUM_GRAY, bold=True, before=0, after=7, line_spacing=1.15)
    _ensure_paragraph_style(doc, "MyWikiTitleDate", alias_of="MyWikiSubtitle")
    _ensure_paragraph_style(doc, "MyWikiSectionTitle", size=layout.section_title_font_size, color=NAVY, bold=True, before=14, after=8, line_spacing=1.18, keep_with_next=True)
    _ensure_paragraph_style(doc, "MyWikiHeading", alias_of="MyWikiSectionTitle")
    _ensure_paragraph_style(doc, "MyWikiIssueLabel", size=layout.issue_label_font_size, color=ACCENT, bold=True, before=14, after=2, line_spacing=1.12, keep_with_next=True)
    _ensure_paragraph_style(doc, "MyWikiIssueTitle", size=layout.issue_title_font_size, color=NAVY, bold=True, before=0, after=6, line_spacing=1.25, keep_with_next=True)
    _ensure_paragraph_style(doc, "MyWikiMetadata", size=layout.metadata_font_size, color=MEDIUM_GRAY, before=0, after=9, line_spacing=1.2, keep_with_next=True)
    _ensure_paragraph_style(doc, "MyWikiMuted", alias_of="MyWikiMetadata")
    _ensure_paragraph_style(doc, "MyWikiSubheading", size=layout.subheading_font_size, color=NAVY, bold=True, before=10, after=5, line_spacing=1.25, keep_with_next=True)
    _ensure_paragraph_style(doc, "MyWikiBody", size=layout.body_font_size, color=DARK_GRAY, before=0, after=7, line_spacing=layout.body_line_spacing)
    _ensure_paragraph_style(doc, "MyWikiBullet", size=layout.bullet_font_size, color=DARK_GRAY, before=0, after=3, line_spacing=layout.bullet_line_spacing)
    _ensure_paragraph_style(doc, "MyWikiSource", size=layout.source_font_size, color=DARK_GRAY, before=0, after=3, line_spacing=layout.source_line_spacing)
    _ensure_paragraph_style(doc, "MyWikiFooter", size=layout.footer_font_size, color=MEDIUM_GRAY, before=0, after=0, line_spacing=1.0)
    _configure_builtin_list_style(doc, "List Bullet", layout, is_numbered=False)
    _configure_builtin_list_style(doc, "List Number", layout, is_numbered=True)


def _ensure_paragraph_style(doc, name: str, *, alias_of: str | None = None, size: float = 10, color: str = DARK_GRAY, bold: bool = False, before: float = 0, after: float = 0, line_spacing: float = 1.0, keep_with_next: bool = False):
    style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if alias_of is not None:
        source = doc.styles[alias_of]
        _apply_style_font(style, size=source.font.size.pt if source.font.size is not None else size, color=str(source.font.color.rgb or color), bold=bool(source.font.bold))
        style.paragraph_format.space_before = source.paragraph_format.space_before
        style.paragraph_format.space_after = source.paragraph_format.space_after
        style.paragraph_format.line_spacing = source.paragraph_format.line_spacing
        style.paragraph_format.keep_with_next = source.paragraph_format.keep_with_next
    else:
        _apply_style_font(style, size=size, color=color, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.keep_with_next = keep_with_next
    style.paragraph_format.widow_control = True
    return style


def _apply_style_font(style, *, size: float, color: str, bold: bool = False) -> None:
    style.font.name = DEFAULT_WORD_FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    _set_style_font_name(style)


def _configure_builtin_list_style(doc, style_name: str, layout: WordLayout, *, is_numbered: bool) -> None:
    style = doc.styles[style_name]
    _apply_style_font(style, size=layout.bullet_font_size, color=DARK_GRAY)
    style.paragraph_format.left_indent = Mm(6.8 if is_numbered else 6.4)
    style.paragraph_format.first_line_indent = Mm(-4.0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = layout.bullet_line_spacing
    style.paragraph_format.widow_control = True


def _configure_header(doc, document: WordReportDocument) -> None:
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        _clear_header_footer(header)
        table = header.add_table(rows=1, cols=2, width=Mm(document.layout.usable_width_mm))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_geometry(table, (document.layout.usable_width_mm * 0.48, document.layout.usable_width_mm * 0.52))
        _clear_table_borders(table)
        _set_row_bottom_border(table.rows[0], color=BORDER_GRAY, size="4")
        left_cell, right_cell = table.rows[0].cells
        for cell in (left_cell, right_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell, top=40, bottom=80, start=0, end=0)
        left_para = left_cell.paragraphs[0]
        _reset_paragraph_spacing(left_para)
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(left_para.add_run(_format_header_date_dots(document.report_date)), size=document.layout.header_date_font_size, color=MEDIUM_GRAY)
        right_para = right_cell.paragraphs[0]
        _reset_paragraph_spacing(right_para)
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(right_para.add_run("MyWiki"), size=document.layout.header_brand_font_size, bold=True, color=NAVY)
        logo_path = _resolve_logo_path()
        if logo_path is not None:
            right_para.add_run(" ")
            right_para.add_run().add_picture(str(logo_path), height=Mm(document.layout.header_logo_height_mm))


def _configure_footer(doc, layout: WordLayout) -> None:
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        _clear_header_footer(footer)
        table = footer.add_table(rows=1, cols=2, width=Mm(layout.usable_width_mm))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_geometry(table, (layout.usable_width_mm * 0.72, layout.usable_width_mm * 0.28))
        _clear_table_borders(table)
        left_cell, right_cell = table.rows[0].cells
        for cell in (left_cell, right_cell):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell, top=40, bottom=0, start=0, end=0)
        left_para = left_cell.paragraphs[0]
        _reset_paragraph_spacing(left_para)
        _set_run_font(left_para.add_run("SK hynix Industry Trend Curation"), size=layout.footer_font_size, color=MEDIUM_GRAY)
        right_para = right_cell.paragraphs[0]
        _reset_paragraph_spacing(right_para)
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_field(right_para, "PAGE", layout.footer_font_size)
        _set_run_font(right_para.add_run(" / "), size=layout.footer_font_size, color=MEDIUM_GRAY)
        _add_field(right_para, "NUMPAGES", layout.footer_font_size)


def _resolve_logo_path() -> Path | None:
    env_value = os.environ.get(WORD_LOGO_ENV)
    if env_value:
        candidate = Path(env_value)
        if candidate.exists():
            return candidate
    if DEFAULT_WORD_LOGO_PATH.exists():
        return DEFAULT_WORD_LOGO_PATH
    return None


def _render_cover_summary(doc, document: WordReportDocument) -> None:
    title = doc.add_paragraph(style="MyWikiReportTitle")
    title.add_run(document.title)
    subtitle = doc.add_paragraph(style="MyWikiSubtitle")
    subtitle.add_run("DAILY INDUSTRY BRIEF")
    meta = doc.add_paragraph(style="MyWikiMetadata")
    meta.add_run(f"기준일 {_format_header_date_dots(document.report_date)} | 생성 {_format_generated_time(document.generated_at)} | v{document.version}")
    _add_section_title(doc, K_EXECUTIVE_SUMMARY)
    if not document.executive_items:
        _add_body_paragraph(doc, K_NO_INFORMATION)
        return
    for index, item in enumerate(document.executive_items, start=1):
        _render_executive_item(doc, index, item, document.layout)


def _render_executive_item(doc, index: int, item: WordExecutiveSummaryLine, layout: WordLayout) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, (12, layout.usable_width_mm - 12))
    _clear_table_borders(table)
    _set_row_bottom_border(table.rows[0], color="E5E7EB", size="4")
    number_cell, content_cell = table.rows[0].cells
    for cell in (number_cell, content_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_margins(cell, top=80, bottom=130, start=60, end=80)
    number_para = number_cell.paragraphs[0]
    _reset_paragraph_spacing(number_para)
    _set_run_font(number_para.add_run(f"{index:02d}"), size=14, bold=True, color=ACCENT)
    title_para = content_cell.paragraphs[0]
    _reset_paragraph_spacing(title_para)
    title_para.paragraph_format.space_after = Pt(3)
    _set_run_font(title_para.add_run(_truncate_text(item.title, 70)), size=11.2, bold=True, color=NAVY)
    summary_para = content_cell.add_paragraph(style="MyWikiBody")
    summary_para.paragraph_format.space_after = Pt(3)
    summary_para.add_run(_truncate_text(item.summary, 145))
    metadata = _format_summary_metadata(item)
    if metadata:
        meta_para = content_cell.add_paragraph(style="MyWikiMetadata")
        meta_para.paragraph_format.space_after = Pt(0)
        meta_para.add_run(metadata)


def _add_section_title(doc, title: str) -> None:
    paragraph = doc.add_paragraph(style="MyWikiSectionTitle")
    paragraph.add_run(title)
    _set_paragraph_bottom_border(paragraph, color=ACCENT, size="6")


def _add_section_gap(doc) -> None:
    paragraph = doc.add_paragraph(style="MyWikiBody")
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)


def _render_issue_section(doc, index: int, section: WordSection) -> None:
    label = doc.add_paragraph(style="MyWikiIssueLabel")
    label.add_run(f"ISSUE {index:02d}")
    title = doc.add_paragraph(style="MyWikiIssueTitle")
    title.add_run(section.title)
    metadata = _format_issue_metadata(section)
    if metadata:
        meta = doc.add_paragraph(style="MyWikiMetadata")
        meta.add_run(metadata)
        _set_paragraph_bottom_border(meta, color=BORDER_GRAY, size="4")
    _add_labeled_paragraph(doc, K_CURRENT_STATUS, section.current_summary)
    _add_bullet_block(doc, K_KEY_FACTS, section.key_facts)
    _add_bullet_block(doc, K_CONTEXT, section.historical_context)
    _add_bullet_block(doc, K_SK_IMPACT, section.implications)
    _add_numbered_block(doc, K_WATCH_POINTS, section.watch_points)


def _render_category_section(doc, document: WordReportDocument) -> None:
    _add_section_title(doc, K_CATEGORY_SUMMARY)
    rows = [document.category_summaries[index : index + 2] for index in range(0, len(document.category_summaries), 2)]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, (document.layout.usable_width_mm / 2, document.layout.usable_width_mm / 2))
    _set_table_borders(table, color=BORDER_GRAY, size="4")
    for row_index, pair in enumerate(rows):
        row = table.rows[row_index]
        _set_row_cant_split(row)
        for cell_index, summary in enumerate(pair):
            cell = row.cells[cell_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_margins(cell, top=140, bottom=140, start=160, end=160)
            _set_cell_shading(cell, LIGHT_GRAY)
            _set_cell_category_content(cell, summary)
        if len(pair) < 2:
            _set_cell_shading(row.cells[1], WHITE)


def _set_cell_category_content(cell, summary: WordCategorySummary) -> None:
    title_para = cell.paragraphs[0]
    title_para.style = "MyWikiSubheading"
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(4)
    title_para.add_run(summary.category)
    for item in (summary.items or (K_NO_MAJOR_TREND,))[:3]:
        paragraph = cell.add_paragraph(style="MyWikiSource")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(_truncate_text(item, 150))


def _render_overall_implications(doc, overall: WordOverallImplications, layout: WordLayout) -> None:
    _add_section_title(doc, K_OVERALL_IMPLICATIONS)
    for label, items, fill in (("OPPORTUNITY", overall.opportunities, "EEF6FF"), ("RISK", overall.risks, "FFF7ED"), ("NEXT WATCH", overall.monitoring_points, "F8FAFC")):
        _render_implication_block(doc, label, items or ("-",), fill, layout)


def _render_implication_block(doc, label: str, items: Sequence[str], fill: str, layout: WordLayout) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, (layout.usable_width_mm,))
    _set_table_borders(table, color=BORDER_GRAY, size="4")
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    _set_cell_shading(cell, fill)
    heading = cell.paragraphs[0]
    heading.style = "MyWikiSubheading"
    heading.paragraph_format.space_before = Pt(0)
    heading.add_run(label)
    for item in items:
        paragraph = cell.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.keep_together = True
        paragraph.add_run(_text_or_dash(item))


def _render_sources_section(doc, sources: tuple[WordSourceLine, ...], layout: WordLayout) -> None:
    _add_section_title(doc, K_SOURCES)
    grouped: dict[str, list[WordSourceLine]] = {}
    for source in sources:
        grouped.setdefault(_source_group_label(source), []).append(source)
    multiple_groups = len(grouped) > 1
    for group_label, rows in grouped.items():
        if multiple_groups:
            heading = doc.add_paragraph(style="MyWikiSubheading")
            heading.add_run(group_label)
        _render_source_table(doc, rows, layout)


def _render_source_table(doc, sources: Sequence[WordSourceLine], layout: WordLayout) -> None:
    widths = (10, 34, layout.usable_width_mm - 90, 26, 20)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_geometry(table, widths)
    _set_table_borders(table, color=BORDER_GRAY, size="4")
    for cell, header in zip(table.rows[0].cells, ("번호", "출처", "제목", "발행일", "링크")):
        _set_cell_shading(cell, LIGHT_GRAY)
        _set_cell_text(cell, header, bold=True, style_name="MyWikiSource")
    _set_row_repeat_header(table.rows[0])
    for index, source in enumerate(sources, start=1):
        row = table.add_row()
        _set_row_cant_split(row)
        cells = row.cells
        _set_cell_text(cells[0], f"{index:02d}", style_name="MyWikiSource", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[1], _source_display_name(source), style_name="MyWikiSource")
        _set_cell_text(cells[2], _text_or_dash(source.title), style_name="MyWikiSource")
        _set_cell_text(cells[3], _text_or_dash(source.published_at), style_name="MyWikiSource", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_source_link_cell(cells[4], source.url)


def _set_source_link_cell(cell, url: str | None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(cell, top=90, bottom=90, start=100, end=100)
    paragraph = cell.paragraphs[0]
    paragraph.style = "MyWikiSource"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if url:
        _add_hyperlink(paragraph, url, K_LINK)
    else:
        paragraph.add_run("-")


def _set_cell_text(cell, text: str, *, bold: bool = False, style_name: str = "MyWikiBody", align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(cell, top=90, bottom=90, start=100, end=100)
    paragraph = cell.paragraphs[0]
    paragraph.style = style_name
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(_text_or_dash(text))
    _set_run_font(run, size=8.8 if style_name == "MyWikiSource" else 10, bold=bold)


def _add_labeled_paragraph(doc, label: str, value: Optional[str]) -> None:
    if not _has_value(value):
        return
    doc.add_paragraph(label, style="MyWikiSubheading")
    doc.add_paragraph(value or "", style="MyWikiBody")


def _add_bullet_block(doc, label: str, items: tuple[str, ...]) -> None:
    if not items:
        return
    doc.add_paragraph(label, style="MyWikiSubheading")
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.keep_together = True


def _add_numbered_block(doc, label: str, items: tuple[str, ...]) -> None:
    if not items:
        return
    doc.add_paragraph(label, style="MyWikiSubheading")
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.keep_together = True


def _add_body_paragraph(doc, text: str):
    return doc.add_paragraph(text, style="MyWikiBody")


def _format_evidence_line(evidence: WordEvidenceLine) -> str:
    base = evidence.evidence_text or K_NO_INFORMATION
    return base if evidence.relevance_score is None else f"{base} (관련도: {evidence.relevance_score:.2f})"


def _format_issue_metadata(section: WordSection) -> str:
    parts = []
    if _has_value(section.category):
        parts.append(section.category)
    if section.importance_score is not None:
        parts.append(f"중요도 {section.importance_score}")
    if _has_value(section.impact_direction):
        parts.append(f"영향 {section.impact_direction}")
    if _has_value(section.time_horizon):
        parts.append(section.time_horizon or "")
    return "   ".join(parts)


def _format_summary_metadata(item: WordExecutiveSummaryLine) -> str:
    parts = []
    if _has_value(item.category):
        parts.append(f"[{item.category}]")
    if item.importance_score is not None:
        parts.append(f"중요도 {item.importance_score}")
    if _has_value(item.impact_direction):
        parts.append(f"영향 {item.impact_direction}")
    if _has_value(item.time_horizon):
        parts.append(item.time_horizon or "")
    return "   ".join(parts)


def _format_generated_time(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return "-"
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        return f"{parsed.hour:02d}:{parsed.minute:02d}"
    except ValueError:
        return raw.split("T", 1)[1][:5] if "T" in raw else raw[:5]


def _format_header_date_dots(value: str) -> str:
    parsed = _parse_date_text(value)
    return value if parsed is None else f"{parsed.year:04d}.{parsed.month:02d}.{parsed.day:02d}"


def _format_header_date_korean(value: str) -> str:
    parsed = _parse_date_text(value)
    return value if parsed is None else f"{parsed.year}년 {parsed.month}월 {parsed.day}일"


def _format_source_date(value: object) -> str:
    if value is None:
        return ""
    raw = str(value).strip()
    return raw[:10].replace("-", ".") if raw else ""


def _parse_date_text(value: str) -> date | None:
    raw = (value or "").strip()
    if not raw:
        return None
    try:
        return date.fromisoformat(raw.split("T", 1)[0].replace(".", "-"))
    except ValueError:
        return None


def _status_value(value: Any) -> str:
    return str(getattr(value, "value", value))


def _enum_text(value: Any) -> str:
    return "" if value is None else normalize_pdf_text(getattr(value, "value", value))


def _clean_optional(value: object) -> str:
    if value is None:
        return ""
    text = normalize_pdf_text(str(value)).strip()
    return "" if text in {"None", "null", K_NO_INFORMATION} else text


def _text_or_dash(value: object) -> str:
    return _clean_optional(value) or "-"


def _has_value(value: object) -> bool:
    return bool(_clean_optional(value))


def _join_title_summary(title: str, summary: str | None) -> str:
    if _has_value(title) and _has_value(summary):
        return f"{title}: {summary}"
    return _text_or_dash(title or summary)


def _truncate_text(text: str, max_chars: int) -> str:
    clean = _text_or_dash(text)
    return clean if len(clean) <= max_chars else clean[: max_chars - 1].rstrip() + "…"


def _source_group_label(source: WordSourceLine) -> str:
    return source.source_type or K_SOURCES


def _source_display_name(source: WordSourceLine) -> str:
    return source.source_name or source.source_type or "-"


def _clear_header_footer(part) -> None:
    for paragraph in list(part.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    for table in list(part.tables):
        tbl = table._element
        tbl.getparent().remove(tbl)


def _reset_paragraph_spacing(paragraph) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0


def _clear_table_borders(table) -> None:
    borders = _get_or_create_table_borders(table)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if border.getparent() is None:
            borders.append(border)
        border.set(qn("w:val"), "nil")


def _set_table_borders(table, *, color: str, size: str) -> None:
    borders = _get_or_create_table_borders(table)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if border.getparent() is None:
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def _get_or_create_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    return borders


def _set_row_bottom_border(row, *, color: str, size: str) -> None:
    for cell in row.cells:
        borders = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            cell._tc.get_or_add_tcPr().append(borders)
        bottom = borders.find(qn("w:bottom")) or OxmlElement("w:bottom")
        if bottom.getparent() is None:
            borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), color)


def _set_paragraph_bottom_border(paragraph, *, color: str, size: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr")) or OxmlElement("w:pBdr")
    if borders.getparent() is None:
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom")) or OxmlElement("w:bottom")
    if bottom.getparent() is None:
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def _set_table_geometry(table, column_widths_mm: Sequence[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_w = tbl.tblPr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl.tblPr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(_mm_to_twips(sum(column_widths_mm))))
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width_mm in column_widths_mm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(_mm_to_twips(width_mm)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for cell, width_mm in zip(row.cells, column_widths_mm):
            cell.width = Mm(width_mm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(_mm_to_twips(width_mm)))


def _set_cell_margins(cell, *, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tc_mar.getparent() is None:
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        margin = tc_mar.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if margin.getparent() is None:
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shading.getparent() is None:
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit")) or OxmlElement("w:cantSplit")
    if cant_split.getparent() is None:
        tr_pr.append(cant_split)


def _set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader")) or OxmlElement("w:tblHeader")
    if tbl_header.getparent() is None:
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _set_style_font_name(style) -> None:
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), DEFAULT_WORD_FONT_NAME)


def _set_run_font(run, *, size: float, bold: bool = False, color: str = DARK_GRAY) -> None:
    run.bold = bold
    run.font.name = DEFAULT_WORD_FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), DEFAULT_WORD_FONT_NAME)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _add_field(paragraph, instruction: str, size: float) -> None:
    begin_run = paragraph.add_run()
    _set_run_font(begin_run, size=size, color=MEDIUM_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instr_run = paragraph.add_run()
    _set_run_font(instr_run, size=size, color=MEDIUM_GRAY)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run._r.append(instr)
    separate_run = paragraph.add_run()
    _set_run_font(separate_run, size=size, color=MEDIUM_GRAY)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    _set_run_font(paragraph.add_run("1"), size=size, color=MEDIUM_GRAY)
    end_run = paragraph.add_run()
    _set_run_font(end_run, size=size, color=MEDIUM_GRAY)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), DEFAULT_WORD_FONT_NAME)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(size)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(r_pr)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _mm_to_twips(value: float) -> int:
    return int(Mm(value).twips)
