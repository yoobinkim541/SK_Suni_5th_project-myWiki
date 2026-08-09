from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from src.analysis.importance_models import ImpactDirection, TimeHorizon
from src.analysis.models import Category
from src.report.models import (
    ReportCategoryGroup,
    ReportCitationDraft,
    ReportExecutiveSummary,
    ReportIssueSummaryRow,
    ReportNewsSource,
    ReportOverallImplications,
    ReportSectionDraft,
    ReportSectionStatus,
    ReportWikiSource,
)
from src.report.word_renderer import (
    DEFAULT_WORD_LOGO_PATH,
    DEFAULT_WORD_TITLE,
    WordReportDocument,
    build_daily_report_word_document,
    build_daily_report_word_filename,
    render_daily_report_word,
)

K_TITLE = "\uc77c\uc77c \uc0b0\uc5c5 \ub3d9\ud5a5 \ub9de\ucda4 \ubcf4\uace0\uc11c"
K_SECTION = "HBM \uc218\uc694 \ud655\ub300\uc640 \uacf5\uae09\ub9dd \ud22c\uc790 \uac00\uc18d"
K_BODY = "\uad6d\ub0b4 \uba54\ubaa8\ub9ac \uc81c\uc870\uc0ac\uc640 \uacf5\uae09\ub9dd \uae30\uc5c5\ub4e4\uc774 AI \uc11c\ubc84 \uc218\uc694 \uc99d\uac00\uc5d0 \ub300\uc751\ud558\uae30 \uc704\ud574 \uc0dd\uc0b0\ub2a5\ub825 \ud655\uc7a5\uacfc \uc7a5\ube44 \ud22c\uc790\ub97c \ub3d9\uc2dc\uc5d0 \ucd94\uc9c4\ud558\uace0 \uc788\uc2b5\ub2c8\ub2e4."
K_EVIDENCE = "HBM \uc218\uc694 \ud655\ub300\uc5d0 \ub9de\ucdb0 \uad00\ub828 \uacf5\uae09\ub9dd \uc804\ubc18\uc5d0\uc11c \ud22c\uc790 \uc18d\ub3c4\uac00 \ube68\ub77c\uc9c0\uace0 \uc788\ub2e4."


def make_section(*, status: ReportSectionStatus = ReportSectionStatus.COMPLETED) -> ReportSectionDraft:
    return ReportSectionDraft(
        issue_key="issue-1",
        representative_analysis_result_id="analysis-1",
        category=Category.PRODUCT_TECHNOLOGY,
        importance_score=88,
        impact_direction=ImpactDirection.OPPORTUNITY,
        time_horizon=TimeHorizon.MID_TERM,
        title=K_SECTION,
        current_summary=K_BODY,
        key_facts=["\ud575\uc2ec \uc0ac\uc2e4 1", "\ud575\uc2ec \uc0ac\uc2e4 2"],
        historical_context=["\uacfc\uac70 \ub9e5\ub77d"],
        implications=["SK\ud558\uc774\ub2c9\uc2a4 \uc218\ud61c\uc0ac \ud655\ub300"],
        watch_points=["\uad00\ucc30 \ud3ec\uc778\ud2b8"],
        news_citations=[
            ReportCitationDraft(
                analysis_result_id="analysis-1",
                document_version_id="doc-ver-1",
                citation_order=1,
                evidence_text=K_EVIDENCE,
                relevance_score=0.92,
            )
        ],
        status=status,
    )


def collect_docx_text(docx_bytes: bytes) -> str:
    document = Document(BytesIO(docx_bytes))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


def test_build_daily_report_word_document_uses_completed_sections_only() -> None:
    completed = make_section(status=ReportSectionStatus.COMPLETED)
    pending = make_section(status=ReportSectionStatus.DRAFTING)

    document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-03",
        version=2,
        sections=[completed, pending],
        generated_at="2026-08-03T09:00:00+09:00",
        report_date=date(2026, 8, 3),
        title=K_TITLE,
    )

    assert document.title == K_TITLE
    assert document.subtitle == "daily-trends-2026-08-03"
    assert document.report_date == "2026-08-03"
    assert document.version == 2
    assert len(document.sections) == 1
    assert document.sections[0].title == K_SECTION


def test_build_daily_report_word_filename_normalizes_spaces() -> None:
    assert build_daily_report_word_filename(report_key="daily trends 2026-08-03", version=3) == "daily-trends-2026-08-03-v3.docx"


def test_render_daily_report_word_returns_docx_bytes() -> None:
    report_document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-03",
        version=1,
        sections=[make_section()],
        generated_at="2026-08-03T09:00:00+09:00",
        report_date=date(2026, 8, 3),
        title=K_TITLE,
    )

    docx_bytes = render_daily_report_word(report_document)
    text = collect_docx_text(docx_bytes)

    assert docx_bytes[:2] == b"PK"
    assert K_TITLE in text
    assert "기준일 2026.08.03" in text
    assert "2026.08.03" in text
    assert "MyWiki" in text
    assert K_SECTION in text
    assert K_BODY in text
    assert K_EVIDENCE in text


def test_render_daily_report_word_preserves_korean_text_in_existing_document() -> None:
    report_document = WordReportDocument(
        title=K_TITLE,
        subtitle="\uc0d8\ud50c \ubcf4\uace0\uc11c",
        generated_at=datetime(2026, 8, 3, 9, 0, 0).isoformat(),
        report_date="2026-08-03",
        version=1,
        sections=build_daily_report_word_document(
            report_key="\uc0d8\ud50c \ubcf4\uace0\uc11c",
            version=1,
            sections=[make_section()],
            generated_at=datetime(2026, 8, 3, 9, 0, 0).isoformat(),
            report_date=date(2026, 8, 3),
            title=K_TITLE,
        ).sections,
    )

    docx_bytes = render_daily_report_word(report_document)
    text = collect_docx_text(docx_bytes)

    assert K_TITLE in text
    assert "\ud575\uc2ec \uc0ac\uc2e4 1" in text
    assert "SK\ud558\uc774\ub2c9\uc2a4 \uc218\ud61c\uc0ac \ud655\ub300" in text


def test_build_daily_report_word_document_defaults_title() -> None:
    document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-03",
        version=1,
        sections=[make_section()],
    )

    assert document.title == DEFAULT_WORD_TITLE


def test_render_daily_report_word_prefers_report_date_over_generated_at() -> None:
    report_document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-03",
        version=1,
        sections=[make_section()],
        generated_at="2026-08-02T23:30:00+00:00",
        report_date=date(2026, 8, 3),
    )

    docx_bytes = render_daily_report_word(report_document)
    text = collect_docx_text(docx_bytes)

    assert "2026.08.03" in text
    assert "기준일 2026.08.03" in text


def _docx_xml(docx_bytes: bytes, name: str) -> str:
    with ZipFile(BytesIO(docx_bytes)) as archive:
        return archive.read(name).decode("utf-8")


def test_render_daily_report_word_embeds_mysuni_logo_to_right_of_mywiki() -> None:
    assert DEFAULT_WORD_LOGO_PATH.exists()
    report_document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-03",
        version=1,
        sections=[make_section()],
        generated_at="2026-08-03T09:00:00+09:00",
        report_date=date(2026, 8, 3),
        title=K_TITLE,
    )

    docx_bytes = render_daily_report_word(report_document)
    header_xml = _docx_xml(docx_bytes, "word/header1.xml")
    header_rels_xml = _docx_xml(docx_bytes, "word/_rels/header1.xml.rels")
    with ZipFile(BytesIO(docx_bytes)) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]

    assert header_xml.index("MyWiki") < header_xml.index("<w:drawing>")
    assert "relationships/image" in header_rels_xml
    assert media_names

def test_render_daily_report_word_uses_executive_brief_sections_and_report_sources() -> None:
    section = make_section()
    report_document = build_daily_report_word_document(
        report_key="daily:ws-1:2026-08-08",
        version=28,
        sections=[section],
        generated_at="2026-08-08T08:27:00+09:00",
        report_date=date(2026, 8, 8),
        title=K_TITLE,
        executive_summaries=[
            ReportExecutiveSummary(
                issue_key="issue-1",
                title="SK하이닉스, AI 메모리 수요 대응 투자 확대",
                summary="HBM·DRAM·NAND 생산능력 확대를 위한 중장기 투자 흐름입니다.",
                importance_score=96,
                impact_direction=ImpactDirection.OPPORTUNITY,
                time_horizon=TimeHorizon.LONG_TERM,
            )
        ],
        issue_summary_rows=[
            ReportIssueSummaryRow(
                issue_key="issue-1",
                category=Category.SUPPLY_PRODUCTION,
                title=section.title,
                importance_score=96,
            )
        ],
        category_groups=[ReportCategoryGroup(category=category, sections=[] if category != Category.SUPPLY_PRODUCTION else [section]) for category in Category],
        overall_implications=ReportOverallImplications(
            opportunities=["HBM 중심의 고부가 메모리 수요 확대를 확인합니다."],
            risks=["장비 수급과 투자 집행 속도 변동성을 점검해야 합니다."],
            monitoring_points=["고객사 수요와 경쟁사 증설 속도를 함께 확인합니다."],
        ),
        news_sources=[
            ReportNewsSource(
                document_version_id="doc-ver-1",
                document_title="매우 긴 기사 제목 " * 12,
                source_name="테스트뉴스",
                published_at="2026-08-07T10:00:00+09:00",
                source_url="https://example.com/very/long/source/url",
            ),
            ReportNewsSource(document_version_id="doc-ver-2"),
        ],
        wiki_sources=[ReportWikiSource(wiki_page_id="wiki-1", wiki_version_id="wiki-v1", wiki_title="HBM 내부 위키")],
    )

    docx_bytes = render_daily_report_word(report_document)
    text = collect_docx_text(docx_bytes)
    document_xml = _docx_xml(docx_bytes, "word/document.xml")
    footer_xml = _docx_xml(docx_bytes, "word/footer1.xml")

    assert "DAILY INDUSTRY BRIEF" in text
    assert "오늘의 핵심 요약" in text
    assert "ISSUE 01" in text
    assert "현재 상황" in text
    assert "카테고리별 정리" in text
    assert "제품·기술" in text
    assert "공급망·생산" in text
    assert "주요 동향 없음" in text
    assert "OPPORTUNITY" in text
    assert "RISK" in text
    assert "NEXT WATCH" in text
    assert "NEWS" in text
    assert "MYWIKI" in text
    assert "https://example.com/very/long/source/url" not in text
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml
    assert "w:cantSplit" in document_xml
    assert "w:tblHeader" in document_xml


def test_render_daily_report_word_sets_layout_styles_for_long_content() -> None:
    section = make_section()
    section.title = "SK하이닉스 " + ("차세대 HBM 투자와 공급망 재편 장기 이슈 " * 8)
    section.current_summary = K_BODY * 25
    section.key_facts = ["한글 bullet 정상 출력: SK하이닉스 용인 청주 공급망 생산 " * 8]
    report_document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-08",
        version=1,
        sections=[section],
        generated_at="2026-08-08T08:27:00+09:00",
        report_date=date(2026, 8, 8),
        news_sources=[ReportNewsSource(document_version_id="doc-ver-1", document_title="긴 출처 제목 " * 20)],
    )

    docx_bytes = render_daily_report_word(report_document)
    document = Document(BytesIO(docx_bytes))
    styles = document.styles
    grids = _docx_xml(docx_bytes, "word/document.xml")

    assert docx_bytes[:2] == b"PK"
    assert styles["MyWikiIssueTitle"].font.size.pt == 14
    assert styles["MyWikiBody"].font.size.pt == 10
    assert styles["MyWikiSubheading"].paragraph_format.keep_with_next is True
    assert styles["MyWikiIssueTitle"].paragraph_format.keep_with_next is True
    assert "w:tblGrid" in grids


def test_render_daily_report_word_handles_empty_categories_sources_and_missing_logo() -> None:
    report_document = build_daily_report_word_document(
        report_key="daily-trends-2026-08-08",
        version=1,
        sections=[make_section()],
        generated_at="2026-08-08T08:27:00+09:00",
        report_date=date(2026, 8, 8),
        category_groups=[ReportCategoryGroup(category=category, sections=[]) for category in Category],
        news_sources=[ReportNewsSource(document_version_id="doc-empty", document_title=None, source_name=None, published_at=None, source_url=None)],
        wiki_sources=[],
    )

    text = collect_docx_text(render_daily_report_word(report_document))

    assert "주요 동향 없음" in text
    assert "doc-empty" in text
    assert "-" in text